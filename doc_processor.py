from docx import Document

def extract_text(path):
    doc = Document(path)
    return "\n".join(para.text for para in doc.paragraphs)

def insert_comments(input_path, output_path, flagged_items):
    doc = Document(input_path)
    for para in doc.paragraphs:
        for phrase in flagged_items:
            if phrase.lower() in para.text.lower():
                para.add_run(f" ← [FLAG: Possibly missing or incomplete: '{phrase}']")
    doc.save(output_path)